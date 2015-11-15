#coding=utf-8
from   ListEditor import *
import pickle
import sys
import os
import copy
<<<<<<< Updated upstream
=======
# import comtypes.client
>>>>>>> Stashed changes
import subprocess
from   PySide.QtGui          import *
from   PySide.QtCore         import *
from   docx		             import Document
from   docx.shared           import Pt
from   docx.enum.style       import WD_STYLE_TYPE
from   time		             import localtime, sleep
from   django.utils.encoding import smart_str, smart_text
from   pyPdf                 import PdfFileReader, PdfFileWriter

class MainWindow(QMainWindow):
	def __init__(self, parent=None):
		super(MainWindow, self).__init__(parent)
		self.setWindowTitle(u'調 假 表 產 生 器')
		self.setWindowIcon(QIcon('./img/changer.png'))
		self.crewList      = './settings/crewlist.pkg'
		self.itemCount	   = 0
		self.exchangeArray = []
		
		self.names         = [u' | --請選擇-- | '] #, u'役男 | 林智煒', u'役男 | 張家榮', u'役男 | 黃凱雋']
		self.listWidget    = QListWidget()
		self.loadCrew()
		
		self.addNewExchange()	
		self.listWidget.setVerticalScrollMode(QAbstractItemView.ScrollPerPixel)
		self.setCentralWidget(self.listWidget)
		self.setupDocker()
		self.setupToolbar()
		self.setupStatusbar()
		# self.showMaximized()
		self.resize(900, 650)

	def loadCrew(self):
		fileName   = self.crewList
		self.names = [u' | --請選擇-- | ']
		stopFlag   = 1
		values     = []
		package    = open( fileName, 'rb' )

		while(stopFlag):
			try:
				values.append( pickle.load( package ))

			except:
				stopFlag = 0
				package.close()

		self.listWidget.clear()
		posArray = [u'請選擇', u'分隊長', u'小隊長', u'隊員', u'役男']
		for value in values:
			elem = u'{0} | {1}'.format(posArray[value[1]], value[2])
			self.names.append(elem)
		package.close()


	def setupStatusbar(self):
		self.progressBar = QProgressBar()
		self.progressBar.setTextVisible(True)
		self.progressBar.setFixedHeight(15)
		self.statusBar().addPermanentWidget(self.progressBar, 1)
		self.progressBar.setValue(0)
		

	def setupToolbar(self):
		self.MainToolbar  = QToolBar(u'工作列') 
		self.MainToolbar.setIconSize(QSize(32,32))

		self.addToolBar( Qt.LeftToolBarArea , self.MainToolbar)
		editAllAction	  = QAction(u'編輯\n清單', self)

		hideAllAction	  = QAction(u'全部\n縮小', self)
		hideAllAction.setIcon(QIcon('./img/minimize.png'))

		showAllAction	  = QAction(u'全部\n展開', self)
		showAllAction.setIcon(QIcon('./img/expand.png'))

		addNewAction	  = QAction(u'新增\n表單', self)
		addNewAction.setIcon(QIcon('./img/addnew.png'))

		retuenAllAction	  = QAction(u'產生\n表單', self)

		editAllAction.triggered.connect  ( self.switchEditDock )
		hideAllAction.triggered.connect  ( self.hideAll        )
		showAllAction.triggered.connect  ( self.showAll        )
		retuenAllAction.triggered.connect( self.setupGenThread )
		addNewAction.triggered.connect   ( self.addNewExchange )

		self.MainToolbar.addAction(editAllAction)
		self.MainToolbar.addAction(hideAllAction)
		self.MainToolbar.addAction(showAllAction)	
		self.MainToolbar.addAction(addNewAction)		
		self.MainToolbar.addAction(retuenAllAction)

	def setupDocker(self):
		self.nameEditor	    = CrewEditor()
		self.nameDockWidget = QDockWidget(u'編輯名單', self)
		self.nameDockWidget.setMinimumWidth(400)
		self.nameDockWidget.setFeatures(QDockWidget.DockWidgetMovable)
		self.nameDockWidget.setAllowedAreas(Qt.LeftDockWidgetArea and Qt.RightDockWidgetArea)
		self.addDockWidget(Qt.RightDockWidgetArea, self.nameDockWidget)
		self.nameDockWidget.setWidget(self.nameEditor)
		self.nameEditor.loadCrew()
		self.nameDockWidget.hide()

	def switchEditDock(self):
		if self.nameDockWidget.isVisible():
			self.nameDockWidget.hide()
		else:
			self.nameDockWidget.show()

	def addNewExchange(self):
		self.itemCount  = self.listWidget.count() + 1

		ClistWidgetItem = QListWidgetItem()
		ClistWidget		= CListWidget(self.listWidget, ClistWidgetItem, self.itemCount, self.names)
		self.exchangeArray.append(ClistWidget)

		ClistWidgetItem.setSizeHint(QSize(565, 300))
		self.listWidget.addItem(ClistWidgetItem)
		ClistWidgetItem.setSizeHint(ClistWidget.sizeHint())
		self.listWidget.addItem(ClistWidgetItem)
		self.listWidget.setItemWidget(ClistWidgetItem, ClistWidget)

	def lockAll(self):
		for index in range(self.listWidget.count()):
			listItem   = self.listWidget.item(index)
			listwidget = self.listWidget.itemWidget (listItem)
			listwidget.lock = True

	def unlockAll(self):
		for index in range(self.listWidget.count()):
			listItem   = self.listWidget.item(index)
			listwidget = self.listWidget.itemWidget (listItem)
			listwidget.lock = False

	def hideAll(self):
		for index in range(self.listWidget.count()):
			listItem   = self.listWidget.item(index)
			listwidget = self.listWidget.itemWidget (listItem)
			listwidget.forceHide()

	def showAll(self):
		for index in range(self.listWidget.count()):
			listItem   = self.listWidget.item(index)
			listwidget = self.listWidget.itemWidget (listItem)
			listwidget.forceShow()	
	
		

	def progress(self, text, progress):
		self.progressBar.setFormat(u'{0} {1}%'.format( text, progress))
		self.progressBar.setValue(progress)
		

	def setupGenThread(self): 
		self.exchangeArray = []
		for index in range(self.listWidget.count()):
			listItem   = self.listWidget.item(index)
			listwidget = self.listWidget.itemWidget (listItem)
			self.exchangeArray.append(listwidget)
		self.yieldThread = yieldSheets(self.exchangeArray)  
		self.yieldThread.threadDone.connect(self.finitializeGenThread, Qt.QueuedConnection)
		self.yieldThread.docxDoneRate.connect(self.progress, Qt.QueuedConnection)
		self.yieldThread.pdfDoneRate.connect(self.progress, Qt.QueuedConnection)
		self.yieldThread.yieldMerged.connect(self.progress, Qt.QueuedConnection)
		self.yieldThread.raiseError.connect(self.progress, Qt.QueuedConnection)
		if not self.yieldThread.isRunning() and self.listWidget.count() != 0:
			self.hideAll()
			self.lockAll()
			self.MainToolbar.setEnabled(False)
			self.yieldThread.start()

	def finitializeGenThread(self, path):
			self.unlockAll()
			self.MainToolbar.setEnabled(True)
			self.openFolder(path)

	def openFolder(self, path):
		if sys.platform == 'linux2':
			def oFolder(path):
				subprocess.check_call(['gnome-open', path])
		elif sys.platform == 'win32':
			def oFolder(path):
				pwd = os.getcwd()
				# subprocess.check_call(['explorer', pwd + '/' + path])
				os.startfile('{0}/output/{1}'.format(pwd, path))
		oFolder(path)

class yieldSheets(QThread):
	threadDone	  = Signal(str)
	pdfDoneRate   = Signal(str, int)
	docxDoneRate  = Signal(str, int)
	yieldMerged   = Signal(str, int)
	raiseError    = Signal(str, int)

	def __init__(self, excArray,  parent=None):
		super(yieldSheets ,self).__init__(parent) 
		self.exchangeArray = excArray

	def run(self):	
		lt             = localtime()
		newfolder      = u'{0}{1}{2}-{3}{4}{5}'.format(lt.tm_year-1911, self.zfill(lt.tm_mon), self.zfill(lt.tm_mday), self.zfill(lt.tm_hour), self.zfill(lt.tm_min), self.zfill(lt.tm_sec))
		c, taskCount   = 0, len(self.exchangeArray)
		template       = Document('./settings/template.docx')
		self.todayText = (u'{0}年{1}月{2}日'.format(lt.tm_year-1911, lt.tm_mon, lt.tm_mday))
		for items in self.exchangeArray:
			c +=2
			doc = copy.deepcopy(template)
			[pos1, pos2, p1I, p2I, date1, date2, reason] = items.returnVal()
			# Dstyle = doc.tables[0].cell(1, 0).paragraphs[0].style
			doc.tables[0].cell(1, 1).text=self.todayText # **date0**
			doc.tables[0].cell(3, 1).paragraphs[0].add_run(pos1)# **receiver**
			doc.tables[0].cell(3, 3).paragraphs[0].add_run(pos2)# **giver**
			doc.tables[0].cell(4, 1).paragraphs[0].add_run(p1I)# **receiver**
			doc.tables[0].cell(4, 3).paragraphs[0].add_run(p2I)# **giver**
			doc.tables[0].cell(5, 0).text=reason# **reason**
			doc.tables[0].cell(9, 9).text=date1
			doc.tables[0].cell(9,11).text=date2
			if not os.path.exists('./output'):
				os.makedirs('./output')
<<<<<<< Updated upstream
			if not os.path.exists('./output/{0}/'.format(newfolder)):	
				os.makedirs('./output/{0}/'.format(newfolder))
			doc.save(str('./output/{1}/exchange{0}.docx'.format(c, newfolder)))
			self.docxDoneRate.emit('generating docx files...' ,int((c/2)*100/taskCount))
		# self.toPDF(   './output/{0}/'.format(newfolder), True)
		# self.mergePDF('./output/{0}/'.format(newfolder), True)
		self.threadDone.emit(newfolder)
=======
			if not os.path.exists('./output/' + newfolder):
				os.makedirs('./output/' + newfolder)
				
			doc.save(str('./output/{1}/exchange{0}.docx'.format(c, newfolder)))
			self.docxDoneRate.emit('generating docx files...' ,int((c/2)*100/taskCount))
		# self.toPDF(   './output/' + newfolder, True)
		# self.mergePDF('./output/' + newfolder, True)

		self.threadDone.emit('./output/' + newfolder )
>>>>>>> Stashed changes

	def zfill(self, num):
		return str(num).zfill(2)

	# def toPDF(self, fileDir, delDocx):
		# docx_files   = [f for f in os.listdir(fileDir) if f.endswith("docx")]
		# wdFormatPDF  =  17
		# try:	
		# 	import comtypes.client
		# 	for filename in docx_files:
		# 		if (u'$' not in filename) and (u'~' not in filename) and (u'½' not in filename):
		# 			comtypes.CoInitialize()
		# 			self.yieldMerged.emit('generating pdf files...' ,int(30))
		# 			in_file  = os.path.abspath(fileDir + filename)
		# 			out_file = os.path.abspath(fileDir + filename.split('.')[0] + '.pdf')
		# 			word     = comtypes.client.CreateObject('Word.Application')
		# 			self.yieldMerged.emit('generating pdf files...' ,int(60))
		# 			doc      = word.Documents.Open(in_file)
		# 			doc.SaveAs(out_file, FileFormat=wdFormatPDF)
		# 			doc.Close()
		# 			word.Quit()
		# 			self.yieldMerged.emit('generating pdf files...' ,int(90))
		# 			if delDocx:
		# 				os.remove(fileDir + filename)
		# 			self.yieldMerged.emit('generating pdf files...' ,int(100))
		# 	return True
		# except:
		# 	try:
		# import win32com.client
		# for filename in docx_files:
		# 	if (u'$' not in filename) and (u'~' not in filename) and (u'½' not in filename):
		# 		win32com.client.pythoncom.CoInitialize()
		# 		self.yieldMerged.emit('generating pdf files...' ,int(30))
		# 		in_file  = os.path.abspath(fileDir + filename)
		# 		out_file = os.path.abspath(fileDir + filename.split('.')[0] + '.pdf')
		# 		word     = win32com.client.Dispatch("Word.Application")
		# 		self.yieldMerged.emit('generating pdf files...' ,int(60))
		# 		doc      = word.Documents.Open(in_file)
		# 		doc.SaveAs(out_file, FileFormat=wdFormatPDF)
		# 		doc.Close()
		# 		word.Quit()
		# 		self.yieldMerged.emit('generating pdf files...' ,int(90))
		# 		if delDocx:
		# 			os.remove(fileDir + filename)
		# 		self.yieldMerged.emit('generating pdf files...' ,int(100))
		# 	# 	return True
			# except:
			# 	return False
			


	def mergePDF(self, fileDir, delPDF):
		output       = PdfFileWriter()
		pdf_files    = [f for f in os.listdir(fileDir) if f.endswith("pdf")]
		c, taskCount = 0, len(pdf_files)
		if pdf_files:
			for filename in pdf_files:
				c += 1
				if (u'$' not in filename) or (u'~' not in filename):
					opened_file = file(fileDir + filename ,"rb")
					self.append_pdf(PdfFileReader(opened_file),output)
					output.write(file(fileDir + "CombinedPages.tmp","wb"))
					opened_file.close()
					self.pdfDoneRate.emit('merging PDF files...', int(c*100/taskCount))
					if  delPDF:
						os.remove(fileDir + filename)
			os.rename(fileDir + 'CombinedPages.tmp', fileDir + 'CombinedPages.pdf')


	def append_pdf(self,input,output):
			[output.addPage(input.getPage(page_num)) for page_num in range(input.numPages)]




class CListWidget(QWidget):
	def __init__(self, host, container, count, names, parent = None):
		super(CListWidget, self).__init__(parent)
		self.state   = True
		self.lock    = False
		self.host    = host
		self.names   = names
		self.number  = count
		self.count   = QLabel('[{0}]'.format(str(count).zfill(2)))
		self.count.setFixedWidth(32)
		self.count.setFixedHeight(32)
		self.count.setAlignment(Qt.AlignCenter)
		self.count.setAutoFillBackground(True)

		self.title   = QLineEdit('Untitled {0}'.format(count))
		self.title.setFixedHeight(32)
		self.title.setReadOnly(True)

		self.hidePB  = QPushButton()
		self.hidePB.setFixedWidth(32)
		self.hidePB.setFixedHeight(32)
		self.hidePB.setIcon(QIcon('./img/minimize.png'))
		self.hidePB.setIconSize(QSize(32,32))
		self.killPB  = QPushButton()
		self.killPB.setFixedWidth(32)
		self.killPB.setFixedHeight(32)
		self.killPB.setIcon(QIcon('./img/kill.png'))
		self.killPB.setIconSize(QSize(32,32))	

		self.container = container
		self.payload = CPayloadWidget(names)
		self.hidePB.clicked.connect(self.hide)
		self.killPB.clicked.connect(self.distory)

		self.vbox1  = QVBoxLayout()
		self.hbox1  = QHBoxLayout()
		self.hbox2  = QHBoxLayout()
		self.hbox3  = QHBoxLayout()
		self.line   = QFrame()
		self.line.setFrameStyle(QFrame.HLine  |  QFrame.Plain)

		self.line.setLineWidth(1)
		self.line.setStyleSheet('border: 1px solid #303030; background-color: #303030;')

		self.hbox1.addWidget(self.count)
		self.hbox1.addWidget(self.title)
		self.hbox1.addWidget(self.hidePB)
		self.hbox1.addWidget(self.killPB)
		self.hbox2.addWidget(self.payload)
		self.hbox3.addWidget(self.line)

		self.vbox1.addLayout(self.hbox1)
		self.vbox1.addLayout(self.hbox2)
		self.vbox1.addLayout(self.hbox3)

		self.setLayout(self.vbox1)


	def hide(self):
		if self.lock == False:	
			self.state = not(self.state)
			self.payload.setVisible(self.state)
			if self.state:
				self.container.setSizeHint(QSize(560, 300))
				self.hidePB.setIcon(QIcon('./img/minimize.png'))
			else:
				self.container.setSizeHint(QSize(560, 50))
				self.hidePB.setIcon(QIcon('./img/expand.png'))
				[pos1, pos2, p1Index, p2Index, p1date, p2date, reason] = self.returnVal()
				self.title.setText(u'{0}  ::  {2}  <----->  {1}  ::  {3}'.format(p1Index, p2Index, p1date.split(u'08時起')[0], p2date.split(u'08時起')[0]))

	def forceHide(self):
		if self.lock == False:
			self.state = True
			self.hide()

	def forceShow(self):
		if self.lock == False:	
			self.state = False
			self.hide()

	def returnVal(self):
		limit   = 5
		p1Index = self.payload.combo1.currentIndex()
		p1      = self.names[p1Index].split(' | ')[1]
		pos1    = self.names[p1Index].split(' | ')[0]

		p2Index = self.payload.combo2.currentIndex()
		p2      = self.names[p2Index].split(' | ')[1]
		pos2    = self.names[p2Index].split(' | ')[0]
		p1date  = self.payload.calendar1.selectedDate()
		p2date  = self.payload.calendar2.selectedDate()

		tmpdate = p1date.addDays(1)
		date1   = u'{0}年{1}月{2}日08時起\n{0}年{3}月{4}日08時止'.format( (p1date.year()-1911), str(p1date.month()).zfill(2), str(p1date.day()).zfill(2), str(tmpdate.month()).zfill(2), str(tmpdate.day()).zfill(2))
		tmpdate = p2date.addDays(1)
		date2   = u'{0}年{1}月{2}日08時起\n{0}年{3}月{4}日08時止'.format( (p2date.year()-1911), str(p2date.month()).zfill(2), str(p2date.day()).zfill(2), str(tmpdate.month()).zfill(2), str(tmpdate.day()).zfill(2))
		reason  = u'事由：職以 {0}年{1}月{2}日 與{5}{6}調休 {0}年{3}月{4}日 一天。'.format( (p1date.year()-1911), str(p1date.month()).zfill(2), str(p1date.day()).zfill(2), str(p2date.month()).zfill(2), str(p2date.day()).zfill(2), pos2, p2)


		return  [pos1, pos2, p1, p2, date1, date2, reason]


	def distory(self):
		for index in xrange(self.number, self.host.count(), 1):
			listItem   = self.host.item(index)
			listwidget = self.host.itemWidget (listItem)
			listwidget.moveUp()
		self.host.takeItem(self.number - 1)

	def moveUp(self):
		self.number -= 1


class CPayloadWidget(QWidget):
	def __init__(self, names, parent = None):
		super(CPayloadWidget, self).__init__(parent)

		self.state   = 0
		self.combo1  = QComboBox()
		self.combo2  = QComboBox()
		self.combo1.setFixedWidth(275)
		self.combo2.setFixedWidth(275)

		self.calendar1 = QCalendarWidget()
		self.calendar2 = QCalendarWidget()
		self.calendar1.showNextMonth()
		self.calendar2.showNextMonth()
		self.calendar1.setVerticalHeaderFormat(QCalendarWidget.NoVerticalHeader)
		self.calendar2.setVerticalHeaderFormat(QCalendarWidget.NoVerticalHeader)
		self.calendar1.setFixedWidth(275)
		self.calendar2.setFixedWidth(275)
		self.calendar1.setFixedHeight(200)
		self.calendar2.setFixedHeight(200)		

		self.hbox1  = QHBoxLayout()
		self.hbox2  = QHBoxLayout()

		self.hbox1.addWidget(self.combo1)
		self.hbox1.addWidget(self.combo2)
		self.combo1.addItems(names)
		self.combo2.addItems(names)

		self.hbox2.addWidget(self.calendar1)
		self.hbox2.addWidget(self.calendar2)

		self.vbox1  = QVBoxLayout()
		self.vbox1.addLayout(self.hbox1)
		self.vbox1.addLayout(self.hbox2)

		self.setLayout(self.vbox1)



app = QApplication(sys.argv)
# MainWindow = listItem()
MainWindow = MainWindow()
MainWindow.show()
def load_stylesheet(pyside=True):
	f = QFile("./settings/style.qss")
	if not f.exists():
		return ""
	else:
		f.open(QFile.ReadOnly  |  QFile.Text)
		ts = QTextStream(f)
		stylesheet = ts.readAll()
		return stylesheet	
app.setStyleSheet(load_stylesheet())
app.exec_()
